from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from office365.runtime.auth.authentication_context import AuthenticationContext
from office365.sharepoint.client_context import ClientContext
import os
import sys
from config import Config
from flask import session
import pythoncom
from app.models.report_soa import ReportSoa

def resource_path(relative_path):
    """L·∫•y ƒë∆∞·ªùng d·∫´n tuy·ªát ƒë·ªëi cho file resources, ho·∫°t ƒë·ªông c·∫£ trong development v√† PyInstaller"""
    try:
        # PyInstaller t·∫°o m·ªôt th∆∞ m·ª•c temp v√† l∆∞u ƒë∆∞·ªùng d·∫´n trong _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class ReportSoaService:
    def __init__(self):
        self.site_url_1 = Config.SHAREPOINT_CONFIG['clearance_flow']['site_url']
        self.site_url_2 = Config.SHAREPOINT_CONFIG.get('clearance', {}).get('site_url', "https://uniconsulting079.sharepoint.com/sites/Clearance")

    def fetch_sharepoint_data(self, from_date, to_date, selected_company):
        print("=== B·∫ÆT ƒê·∫¶U KI·ªÇM TRA D·ªÆ LI·ªÜU ===")

        # Ki·ªÉm tra x√°c th·ª±c
        if 'user' not in session or not session['user'].get('authenticated'):
            print("Debug - User not authenticated")
            return {'error': 'Vui l√≤ng ƒëƒÉng nh·∫≠p l·∫°i'}
        
        username = session['user'].get('username')
        password = session['user'].get('password')
        
        if not username or not password:
            print("Debug - Missing credentials")
            return {'error': 'Vui l√≤ng ƒëƒÉng nh·∫≠p l·∫°i'}

        try:
            # Chuy·ªÉn ƒë·ªãnh d·∫°ng ng√†y
            from_date_dt = datetime.strptime(from_date, "%d/%m/%Y")
            to_date_dt = datetime.strptime(to_date, "%d/%m/%Y")
            from_date_str = from_date_dt.strftime("%m/%d/%Y")
            to_date_str = to_date_dt.strftime("%m/%d/%Y")

            # X√°c th·ª±c v·ªõi SharePoint
            try:
                ctx1 = self._get_authenticated_context(self.site_url_1, username, password)
                ctx2 = self._get_authenticated_context(self.site_url_2, username, password)
            except Exception as e:
                print(f"Debug - SharePoint authentication error: {str(e)}")
                return {'error': 'L·ªói x√°c th·ª±c SharePoint'}

            # L·∫•y d·ªØ li·ªáu t·ª´ Debit/Credit Note
            debit_credit_dict = self._get_debit_credit_data(ctx1)
            
            # L·∫•y d·ªØ li·ªáu t·ª´ Freight BL
            filtered_data = self._get_freight_bl_data(ctx2, from_date_dt, to_date_dt, selected_company, debit_credit_dict)

            company_info = self._get_company_info(filtered_data)
            print("=== K·∫æT TH√öC KI·ªÇM TRA ===")
            
            return filtered_data, company_info

        except Exception as e:
            print(f"Debug - Error in fetch_sharepoint_data: {str(e)}")
            return {'error': f'L·ªói khi l·∫•y d·ªØ li·ªáu: {str(e)}'}

    def _get_authenticated_context(self, site_url, username, password):
        """T·∫°o context ƒë√£ x√°c th·ª±c cho SharePoint"""
        try:
            auth_context = AuthenticationContext(site_url)
            if auth_context.acquire_token_for_user(username, password):
                ctx = ClientContext(site_url, auth_context)
                # Verify authentication
                ctx.load(ctx.web)
                ctx.execute_query()
                return ctx
            raise Exception("Kh√¥ng th·ªÉ x√°c th·ª±c v·ªõi SharePoint")
        except Exception as e:
            print(f"Debug - Authentication error: {str(e)}")
            raise Exception("L·ªói x√°c th·ª±c SharePoint")

    def _get_debit_credit_data(self, ctx):
        list_debit_credit = ctx.web.lists.get_by_title("Debit/Credit Note")
        debit_credit_items = []
        paged_items = list_debit_credit.items.get().top(1000).execute_query()
        debit_credit_items.extend(paged_items)
        
        while True:
            try:
                paged_items = list_debit_credit.items.get_next_page(paged_items)
                ctx.load(paged_items)
                ctx.execute_query()
                debit_credit_items.extend(paged_items)
            except Exception:
                break

        debit_credit_dict = {}
        for item in debit_credit_items:
            id_of_list = item.properties.get("IDofList")
            if id_of_list is not None:
                if id_of_list not in debit_credit_dict:
                    debit_credit_dict[id_of_list] = []
                debit_credit_dict[id_of_list].append({
                    "DEBIT": item.properties.get("DEBIT", None),
                    "CREDIT": item.properties.get("CREDIT", None)
                })
        return debit_credit_dict

    def _get_freight_bl_data(self, ctx, from_date_dt, to_date_dt, selected_company, debit_credit_dict):
        list_clearance = ctx.web.lists.get_by_title("Freight BL")
        freight_bl_items = []
        paged_items = list_clearance.items.get().top(1000).execute_query()
        freight_bl_items.extend(paged_items)
        
        while True:
            try:
                paged_items = list_clearance.items.get_next_page(paged_items)
                ctx.load(paged_items)
                ctx.execute_query()
                freight_bl_items.extend(paged_items)
            except Exception:
                break

        filtered_data = []
        for item in freight_bl_items:
            on_board_date_str = item.properties.get("ON_x002d_BOARDDATE")
            if not on_board_date_str:
                continue

            try:
                on_board_date = datetime.strptime(on_board_date_str.split("T")[0], "%Y-%m-%d")
            except ValueError:
                continue

            if not (from_date_dt <= on_board_date <= to_date_dt):
                continue

            apply_to = item.properties.get("APPLYTO", "")
            if selected_company.lower() not in apply_to.lower():
                continue

            clearance_id = item.properties.get("ID")
            title = item.properties.get("Title", "").strip()

            if clearance_id in debit_credit_dict:
                total_debit = 0
                total_credit = 0
                for debit_credit in debit_credit_dict[clearance_id]:
                    try:
                        debit = float(debit_credit["DEBIT"]) if debit_credit["DEBIT"] is not None else 0
                        credit = float(debit_credit["CREDIT"]) if debit_credit["CREDIT"] is not None else 0
                    except (ValueError, TypeError):
                        debit = 0
                        credit = 0
                    total_debit += debit
                    total_credit += credit

                statement = ReportSoa(
                    hbl_no=title,
                    mbl=item.properties.get("MBL", ""),
                    date=on_board_date.strftime("%d/%m/%Y"),
                    debit=total_debit if total_debit != 0 else None,
                    credit=total_credit if total_credit != 0 else None,
                    apply_to=apply_to
                )
                filtered_data.append(statement.to_dict())

        return filtered_data

    def _get_company_info(self, filtered_data):
        """L·∫•y th√¥ng tin c√¥ng ty t·ª´ d·ªØ li·ªáu ƒë√£ l·ªçc"""
        if filtered_data:
            # L·∫•y APPLYTO t·ª´ b·∫£n ghi ƒë·∫ßu ti√™n
            first_record = filtered_data[0]
            return first_record.get("apply_to", "")
        return ""

    def create_word_file(self, data, company_name, company_info, from_date, to_date, save_dir, file_format='docx'):
        """T·∫°o file Word ho·∫∑c PDF d·ª±a tr√™n file_format"""
        temp_docx = None
        final_filename = None
        
        try:
            # T·∫°o document Word
            doc = Document()
            
            # T·∫°o n·ªôi dung document
            self._create_header(doc)
            self._create_title_section(doc, company_info, from_date, to_date)
            self._create_data_table(doc, data)
            self._create_footer(doc)
            
            base_name = f"{company_name}_report"
            
            if file_format.lower() == 'pdf':
                print("üîÑ T·∫°o file PDF...")
                return self._convert_to_pdf(doc, base_name, save_dir)
            else:
                print("üìÑ T·∫°o file Word...")
                # T·∫°o file Word
                final_filename = f"{base_name}.docx"
                doc.save(os.path.join(save_dir, final_filename))
                print(f"‚úÖ ƒê√£ t·∫°o file Word: {final_filename}")
                return final_filename
            
        except Exception as e:
            # ƒê·∫£m b·∫£o x√≥a file t·∫°m n·∫øu c√≥ l·ªói
            if temp_docx and os.path.exists(temp_docx):
                try:
                    os.remove(temp_docx)
                except:
                    pass
            raise Exception(f"Error creating document: {str(e)}")

    def _convert_to_pdf(self, doc, base_name, save_dir):
        """Convert Word document to PDF v·ªõi multiple fallback options"""
        import uuid
        temp_id = str(uuid.uuid4())
        temp_docx = os.path.join(save_dir, f"temp_{temp_id}.docx")
        
        try:
            # L∆∞u file Word t·∫°m th·ªùi
            doc.save(temp_docx)
            print(f"üìÑ ƒê√£ t·∫°o file Word t·∫°m: {temp_docx}")
            
            final_filename = f"{base_name}.pdf"
            pdf_path = os.path.join(save_dir, final_filename)
            
            # Option 1: docx2pdf (works on Linux with LibreOffice)
            success = self._try_docx2pdf_conversion(temp_docx, pdf_path)
            
            if not success:
                # Option 2: pandoc (n·∫øu c√≥ c√†i)
                success = self._try_pandoc_conversion(temp_docx, pdf_path)
            
            if not success:
                # Option 3: LibreOffice headless (n·∫øu c√≥ c√†i)
                success = self._try_libreoffice_conversion(temp_docx, pdf_path)
            
            if success:
                print(f"‚úÖ ƒê√£ chuy·ªÉn ƒë·ªïi th√†nh c√¥ng sang PDF: {final_filename}")
                return final_filename
            else:
                # Fallback: Tr·∫£ v·ªÅ file Word n·∫øu kh√¥ng convert ƒë∆∞·ª£c
                print("‚ö†Ô∏è Kh√¥ng th·ªÉ chuy·ªÉn ƒë·ªïi sang PDF, tr·∫£ v·ªÅ file Word")
                fallback_filename = f"{base_name}.docx"
                fallback_path = os.path.join(save_dir, fallback_filename)
                
                # Copy file t·∫°m th√†nh file cu·ªëi c√πng
                import shutil
                shutil.copy2(temp_docx, fallback_path)
                return fallback_filename
                
        except Exception as e:
            print(f"‚ùå L·ªói khi chuy·ªÉn ƒë·ªïi PDF: {str(e)}")
            # Fallback: Tr·∫£ v·ªÅ file Word
            fallback_filename = f"{base_name}.docx"
            fallback_path = os.path.join(save_dir, fallback_filename)
            
            try:
                import shutil
                shutil.copy2(temp_docx, fallback_path)
                return fallback_filename
            except:
                raise Exception("Kh√¥ng th·ªÉ t·∫°o file")
                
        finally:
            # X√≥a file t·∫°m
            if temp_docx and os.path.exists(temp_docx):
                try:
                    os.remove(temp_docx)
                    print(f"üóëÔ∏è ƒê√£ x√≥a file t·∫°m: {temp_docx}")
                except Exception as e:
                    print(f"‚ö†Ô∏è Kh√¥ng th·ªÉ x√≥a file t·∫°m: {str(e)}")

    def _try_docx2pdf_conversion(self, docx_path, pdf_path):
        """Th·ª≠ convert b·∫±ng docx2pdf"""
        try:
            from docx2pdf import convert # type: ignore
            print("üîÑ ƒêang th·ª≠ docx2pdf conversion...")
            convert(docx_path, pdf_path)
            
            if os.path.exists(pdf_path):
                print("‚úÖ docx2pdf conversion th√†nh c√¥ng!")
                return True
            return False
            
        except ImportError:
            print("‚ö†Ô∏è docx2pdf kh√¥ng ƒë∆∞·ª£c c√†i ƒë·∫∑t")
            return False
        except Exception as e:
            print(f"‚ùå docx2pdf conversion th·∫•t b·∫°i: {str(e)}")
            return False

    def _try_pandoc_conversion(self, docx_path, pdf_path):
        """Th·ª≠ convert b·∫±ng pandoc"""
        try:
            import subprocess
            print("üîÑ ƒêang th·ª≠ pandoc conversion...")
            
            result = subprocess.run([
                'pandoc', docx_path, '-o', pdf_path,
                '--pdf-engine=xelatex'
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0 and os.path.exists(pdf_path):
                print("‚úÖ pandoc conversion th√†nh c√¥ng!")
                return True
            else:
                print(f"‚ùå pandoc conversion th·∫•t b·∫°i: {result.stderr}")
                return False
                
        except FileNotFoundError:
            print("‚ö†Ô∏è pandoc kh√¥ng ƒë∆∞·ª£c c√†i ƒë·∫∑t")
            return False
        except Exception as e:
            print(f"‚ùå pandoc conversion th·∫•t b·∫°i: {str(e)}")
            return False

    def _try_libreoffice_conversion(self, docx_path, pdf_path):
        """Th·ª≠ convert b·∫±ng LibreOffice headless"""
        try:
            import subprocess
            print("üîÑ ƒêang th·ª≠ LibreOffice conversion...")
            
            # Get output directory
            output_dir = os.path.dirname(pdf_path)
            
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', output_dir, docx_path
            ], capture_output=True, text=True, timeout=120)
            
            # LibreOffice t·∫°o file v·ªõi t√™n g·ªëc
            expected_pdf = os.path.join(output_dir, 
                                     os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
            
            if os.path.exists(expected_pdf):
                # Rename n·∫øu c·∫ßn
                if expected_pdf != pdf_path:
                    import shutil
                    shutil.move(expected_pdf, pdf_path)
                print("‚úÖ LibreOffice conversion th√†nh c√¥ng!")
                return True
            else:
                print(f"‚ùå LibreOffice conversion th·∫•t b·∫°i: {result.stderr}")
                return False
                
        except FileNotFoundError:
            print("‚ö†Ô∏è LibreOffice kh√¥ng ƒë∆∞·ª£c c√†i ƒë·∫∑t")
            return False
        except Exception as e:
            print(f"‚ùå LibreOffice conversion th·∫•t b·∫°i: {str(e)}")
            return False

    def _create_header(self, doc):
        # Thi·∫øt l·∫≠p ƒë·ªô r·ªông trang v√† margin
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # T·∫°o b·∫£ng header v·ªõi 2 c·ªôt
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        
        widths = [0.5, 2.0, 1.8, 0.8, 0.6, 1.5, 1.5]  # ƒê·ªô r·ªông c√°c c·ªôt c·ªßa b·∫£ng d·ªØ li·ªáu
        total_width = sum(widths)  # = 8.7 inches
        
        
        table.columns[0].width = Inches(5.7)  # ƒêi·ªÅu ch·ªânh c·ªôt tr√°i
        table.columns[1].width = Inches(3.0)  # ƒêi·ªÅu ch·ªânh c·ªôt ph·∫£i
        
        # Th√¥ng tin c√¥ng ty b√™n tr√°i
        left_cell = table.cell(0, 0)
        company_infos = [
            'UNI CONSULTING CO., LTD',
            '113A, STREET 109, QUARTER 5,',
            'PHUOC LONG B WARD, THU DUC',
            'CITY, HCMC, VIETNAM.',
            'TEL: +84 028 36365922'
        ]
        for line in company_infos:
            p = left_cell.add_paragraph()
            p.paragraph_format.line_spacing = 0.8
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0)
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
            run.font.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Logo b√™n ph·∫£i
        right_cell = table.cell(0, 1)
        p = right_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.right_indent = Inches(0)
        run = p.add_run()
        try:
            logo_path = os.path.join("app", "static", "images", "logo.png")
            if not os.path.exists(logo_path):
                logo_path = resource_path(os.path.join("app", "static", "images", "logo.png"))
            run.add_picture(logo_path, width=Inches(1.2))
        except Exception as e:
            print(f"Warning: Could not load logo - {str(e)}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Th√™m kho·∫£ng c√°ch sau b·∫£ng
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def _create_title_section(self, doc, company_info, from_date, to_date):
        # Ti√™u ƒë·ªÅ
        title = doc.add_paragraph()
        title.paragraph_format.line_spacing = 1.0
        title.paragraph_format.space_after = Pt(12)
        title_run = title.add_run('STATEMENT OF ACCOUNT')
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = 'Times New Roman'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Ng√†y hi·ªán t·∫°i
        date = doc.add_paragraph()
        date.paragraph_format.line_spacing = 1.0
        date.paragraph_format.space_after = Pt(12)
        current_date = datetime.now().strftime("%d/%m/%Y")
        date_run = date.add_run(current_date)
        date_run.font.size = Pt(10)
        date_run.font.name = 'Times New Roman'
        date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Th√¥ng tin c√¥ng ty t·ª´ SharePoint
        if company_info:
            p_company = doc.add_paragraph()
            p_company.paragraph_format.line_spacing = 1.0
            p_company.paragraph_format.space_after = Pt(12)
            
            # Th√™m "To: " v√†o tr∆∞·ªõc th√¥ng tin c√¥ng ty
            company_run = p_company.add_run(f"To: {company_info}")
            company_run.bold = True
            company_run.font.size = Pt(10)
            company_run.font.name = 'Times New Roman'

        # Th√™m kho·∫£ng c√°ch tr∆∞·ªõc th√¥ng tin th√™m
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # Th√¥ng tin th√™m
        info_para = doc.add_paragraph()
        info_para.paragraph_format.line_spacing = 1.0
        info_para.paragraph_format.space_after = Pt(6)
        info_para.paragraph_format.left_indent = Inches(0)
        info_para.paragraph_format.right_indent = Inches(0)
        
        info_text = f"SEA/AIR(ALL)         POL/POD:ALL\nPERIOD: {from_date} ~ {to_date}"
        info_run = info_para.add_run(info_text)
        info_run.bold = True
        info_run.font.size = Pt(10)
        info_run.font.name = 'Times New Roman'

    def _create_data_table(self, doc, data):
        # S·∫Øp x·∫øp d·ªØ li·ªáu
        sorted_data = sorted(data, key=lambda x: x["HBL_NO"])
        
        # T·∫°o b·∫£ng
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        table.allow_autofit = True
        
        # Thi·∫øt l·∫≠p ƒë·ªô r·ªông c·ªôt m·ªõi
        widths = [0.5, 1.5, 1.8, 0.8, 0.6, 1.7, 1.7]  # T·ªïng = 8.7 inches
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
        
        # Header
        headers = ["SEQ", "H.B/L NO.", "M.BL(AWB) NO.", "DATE", "CURR", "DEBIT", "CREDIT"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # D·ªØ li·ªáu
        total_debit = 0
        total_credit = 0
        
        for idx, item in enumerate(sorted_data, start=1):
            row_cells = table.add_row().cells
            
            debit = item.get("DEBIT", 0)
            credit = item.get("CREDIT", 0)
            
            try:
                debit_value = float(debit) if debit is not None else 0
                credit_value = float(credit) if credit is not None else 0
            except (ValueError, TypeError):
                debit_value = 0
                credit_value = 0
            
            total_debit += debit_value
            total_credit += credit_value
            
            # ƒêi·ªÅn d·ªØ li·ªáu v√†o t·ª´ng √¥
            row_cells[0].text = str(idx)
            row_cells[1].text = item["HBL_NO"]
            row_cells[2].text = item["MBL"]
            row_cells[3].text = item["Date"]
            row_cells[4].text = "USD"
            row_cells[5].text = f"{debit_value:,.2f}" if debit_value != 0 else ""
            row_cells[6].text = f"{credit_value:,.2f}" if credit_value != 0 else ""
            
            # Format font v√† cƒÉn ch·ªânh
            for i, cell in enumerate(row_cells):
                cell.paragraphs[0].runs[0].font.size = Pt(11)
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                # CƒÉn gi·ªØa cho c·ªôt SEQ, DATE, CURR
                if i in [0, 3, 4,1,2]:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # CƒÉn ph·∫£i cho c·ªôt DEBIT, CREDIT
                elif i in [5, 6]:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # CƒÉn tr√°i cho c√°c c·ªôt c√≤n l·∫°i
                else:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Th√™m d√≤ng t·ªïng
        self._add_total_row(table, total_debit, total_credit)
        
        # Th√™m d√≤ng ch√™nh l·ªách
        self._add_difference_row(table, total_debit, total_credit)

        # NgƒÉn kh√¥ng cho b·∫£ng b·ªã t√°ch trang
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT.AT_LEAST
            row.height = Pt(20)  # Chi·ªÅu cao t·ªëi thi·ªÉu c·ªßa m·ªói d√≤ng
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1

    def _add_total_row(self, table, total_debit, total_credit): 
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[3])
        row_cells[0].text = "SUB TOTAL"
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[4].text = "USD"
        row_cells[5].text = f"{total_debit:,.2f}"
        row_cells[6].text = f"{total_credit:,.2f}"
        
        for cell in [row_cells[0], row_cells[4], row_cells[5], row_cells[6]]:
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].bold = True

    def _add_difference_row(self, table, total_debit, total_credit): 
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[3])
        row_cells[0].text = "TOTAL"
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row_cells[4].text = "USD"
        difference = total_debit - total_credit
        row_cells[5].text = f"{difference:,.2f}"
        row_cells[6].text = ""
        
        for cell in [row_cells[0], row_cells[4], row_cells[5]]:
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].bold = True

    def _create_footer(self, doc):
        # Th√™m kho·∫£ng c√°ch tr∆∞·ªõc footer
        for _ in range(4):  # Th√™m 3 d√≤ng tr·ªëng + 1 d√≤ng space ban ƒë·∫ßu
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(6)
            p_space.paragraph_format.space_after = Pt(6)

        # Th√™m m·ªôt g·∫°ch ngang ƒë∆°n
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(6)
        p_line.paragraph_format.line_spacing = 1
        
        pBdr = OxmlElement('w:pBdr')
        b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '1')
        b.set(qn('w:color'), '000000')
        pBdr.append(b)
        p_line._p.get_or_add_pPr().append(pBdr)

        # Th√™m n·ªôi dung footer
        warning = doc.add_paragraph()
        warning.add_run('PLS CONTACT US WITHIN 3 DAYS IF THERE ARE ANY DISCREPANCIES.').font.size = Pt(10)
        
        payment_title = doc.add_paragraph()
        payment_title.add_run('PAYMENT TO:').font.size = Pt(10)
        payment_title.runs[0].bold = True
        
        # Th√¥ng tin ng√¢n h√†ng trong khung
        bank_info = [
            ('BENEFICIARY: UNI CONSULTING CO., LTD', True),
            ('BANK ADDRESS: SHINHAN BANK VIETNAM LIMITED - BIEN HOA BRANCH\n' +
             'Sonadezi Building, 9th floor, No. 01, Street 1, Bien Hoa 1 industrial zone An\n' +
             'Binh Ward, Bien Hoa City, Dong Nai Province, Viet Nam', True),
            ('A/C: 700-019-156046 (USD)\nSWIFT CODE: SHBKVNXXXXX', True)
        ]
        
        for text, add_border in bank_info:
            p = doc.add_paragraph()
            p.add_run(text).font.size = Pt(10)
            if add_border:
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                for border in ['top', 'bottom', 'left', 'right']:
                    b = OxmlElement(f'w:{border}')
                    b.set(qn('w:val'), 'single')
                    b.set(qn('w:sz'), '6')
                    b.set(qn('w:space'), '1')
                    b.set(qn('w:color'), 'auto')
                    pBdr.append(b)
                pPr.append(pBdr) 